import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

def set_cell_border(cell, **kwargs):
    """
    Set cell's border. This function is used to hide the table borders by setting them to zero size.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Create a 'w:tcBorders' element for the cell
    tcBorders = OxmlElement('w:tcBorders')
    for key, value in kwargs.items():
        tag = 'w:{}'.format(key)
        element = OxmlElement(tag)
        for k, v in value.items():
            # Ensure the value is converted to string
            element.set(qn('w:{}'.format(k)), str(v))  # Convert integer values to strings
        tcBorders.append(element)
    
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=None, bottom=None, left=None, right=None):
    """
    Set the margins of a table cell.

    :param cell: The cell to adjust the margins for.
    :param top: Top margin in Pt.
    :param bottom: Bottom margin in Pt.
    :param left: Left margin in Pt.
    :param right: Right margin in Pt.
    """
    cell_margins = OxmlElement("w:tcMar")

    if top is not None:
        top_margin = OxmlElement('w:top')
        top_margin.set(qn('w:w'), str(top))
        top_margin.set(qn('w:type'), 'dxa')
        cell_margins.append(top_margin)

    if bottom is not None:
        bottom_margin = OxmlElement('w:bottom')
        bottom_margin.set(qn('w:w'), str(bottom))
        bottom_margin.set(qn('w:type'), 'dxa')
        cell_margins.append(bottom_margin)

    if left is not None:
        left_margin = OxmlElement('w:left')
        left_margin.set(qn('w:w'), str(left))
        left_margin.set(qn('w:type'), 'dxa')
        cell_margins.append(left_margin)

    if right is not None:
        right_margin = OxmlElement('w:right')
        right_margin.set(qn('w:w'), str(right))
        right_margin.set(qn('w:type'), 'dxa')
        cell_margins.append(right_margin)

    cell._tc.get_or_add_tcPr().append(cell_margins)

def add_header(doc, name, address, email, phone):
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(30)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Adjust paragraph spacing (before and after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Insert Line Below
    p_border = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '8')  # Border size, 4 = 1/2 pt
    bottom_border.set(qn('w:space'), '1')  # Space between text and border
    bottom_border.set(qn('w:color'), '000000')  # Border color, 'auto' or hex color code
    p_border.append(bottom_border)
    p._element.get_or_add_pPr().append(p_border)

    p = doc.add_paragraph()
    p.add_run(address + '\t\t').bold = False
    p.add_run(email + '\t\t').italic = True
    p.add_run(phone).italic = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

def add_section_title(doc, title):
    # Add Title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Adjust paragraph spacing (before and after)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)

    # Insert Line Below
    p_border = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '8')  # Border size, 4 = 1/2 pt
    bottom_border.set(qn('w:space'), '1')  # Space between text and border
    bottom_border.set(qn('w:color'), '000000')  # Border color, 'auto' or hex color code
    p_border.append(bottom_border)
    p._element.get_or_add_pPr().append(p_border)

def add_education(doc, school, degree, major, graduation_date):
    p = doc.add_paragraph()
    run = p.add_run(f"{degree} in {major}, {school}, {graduation_date}")
    run.bold = True

def add_experience(doc, 
                   title: str = '', 
                   company: str = '',
                   dates: str = '', 
                   location: str = '', 
                   bullets: list = []):
    
    # Add a table with one row and two columns
    table = doc.add_table(rows=2, cols=2)
    
    for (row, side, text) in [
        (0, 0, title),
        (1, 0, company),
        (0, 1, dates),
        (1, 1, location)
    ]:
        cell = table.cell(row, side)
        set_cell_border(cell, bottom={"sz": 0, "val": "nil", "color": "FFFFFF"})
        set_cell_margins(cell, top=30, bottom=30, left=30, right=30)
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        

        paragraph = cell.paragraphs[0]
        if (side == 0): paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else: paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        if (row == 0 and side == 0): paragraph.runs[0].bold = True
        if (row == 1 and side == 0): paragraph.runs[0].italic = True

        # Set the height for the row
        row = table.rows[row]
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(Pt(50)))  # Set the row height to 25 points
        trHeight.set(qn('w:hRule'), "atLeast")  # Use "atLeast" or "exact" for fixed height
        trPr.append(trHeight)

    # Adjust the table to span the width of the document
    table.autofit = False
    table.columns[0].width = Inches(4)
    table.columns[1].width = Inches(4)


    for bullet in bullets:
        p = doc.add_paragraph(f"  {bullet}", style='List Bullet')
        p.paragraph_format.left_indent = Pt(25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

def add_project(doc, title, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title)
    run.bold = True
    for bullet in bullets:
        p = doc.add_paragraph(f"  {bullet}", style='List Bullet')
        p.paragraph_format.left_indent = Pt(25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

def rank(df_exp, df_proj, budget = 30):
    # rank first vby 'keyword_count' and then for ties, rank by 'embedding' similarity
    df_exp['type'] = 'exp'
    df_proj['type'] = 'proj'
    df = pd.concat([df_exp, df_proj], ignore_index=True)
    sorted_df = df.sort_values(by=['keyword_count', 'similarity'], ascending=[False, False])

    unique_tags = set()
    selected_exp = []
    selected_proj = []
    total_cost = 0

    for _, row in sorted_df.iterrows():
        row_cost = 1.5 + 2*(row['title'] not in unique_tags)
        if total_cost + row_cost <= budget:
            total_cost += row_cost
            unique_tags.add(row['title'])
            if row['type'] == 'exp':
                selected_exp.append(row)
            else:
                selected_proj.append(row)
        else:
            break   

    return selected_exp, selected_proj

def make_resume(df_exp, df_proj, username):
    doc = Document()

    section = doc.sections[0]

    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)

    add_header(doc, 'John Doe', '1234 Elm St, Springfield, IL 62701', 'asdf@gmail.com', '123-456-7890')

    add_section_title(doc, 'Education')
    add_education(doc, 'University of Illinois Urbana-Champaign', 'Bachelor of Science', 'Computer Science', 'May 2020')

    exp, proj = rank(df_exp, df_proj)
    print(exp)

    add_section_title(doc, 'Experience')
    for row in exp:
        print(f'Adding experience: {row["title"]}: {row["bullet"]}')
        # add_experience(doc, row['title'], row['company'], row['dates'], bullets=row['bullets'])

    # add_section_title(doc, 'Projects')
    # for _, row in proj.iterrows():
    #     add_project(doc, row['title'], row['bullets'])

    filename = f'./pdfs/{username}_resume.docx'
    doc.save(filename)
    return filename